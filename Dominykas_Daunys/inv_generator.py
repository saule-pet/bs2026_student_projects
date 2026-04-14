# %% Bibliotekų diegimas

# Sistemos nustatymai
import locale
import os
from dotenv import load_dotenv

# Laikas
from datetime import datetime, timedelta
from collections import defaultdict

# Kalendoriaus duomenys
import caldav
from icalendar import Calendar

# Dokumentų (word, pdf) gamyba
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import num2words
import comtypes.client

# Emailų siuntimas
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders

# %% Iniciacija

# Sukuriame aplanką sąskaitoms, jeigu jo dar nėra
os.makedirs('Saskaitos_out', exist_ok=True)

# Užtikrinu, kad laikas - Lietuvos
locale.setlocale(locale.LC_TIME, 'lt_LT.UTF-8')

# Laikotarpio pradžia ir pabaiga
start_date = datetime(2026, 3, 1)
end_date = datetime(2026, 4, 15)

# Užkraunu duomenis iš env failo
load_dotenv()

# Gmail credentials
gmail_user = os.getenv('GMAIL_USER')
gmail_password = os.getenv('GMAIL_PASSWORD')

# iCloud credentials ir CalDAV URL
username = os.getenv('ICLOUD_USER')
password = os.getenv('ICLOUD_PASSWORD')
url = 'https://caldav.icloud.com/'

# Prisijungiu prie icloud
try:
    client = caldav.DAVClient(url, username=username, password=password)
    principal = client.principal() # gauk duomenis apie vartotoją
    calendars = principal.calendars() # gauk vartotojo kalendorius
except Exception as e: #apsidraudimui, jeigu neconectintu su icloudu
    print(f"Error connecting to iCloud: {e}")
    calendars = []

# %% Duomenys

# Sukuriu tuščius žodynus, į kuriuos sudėsiu pamokų trukmę
durations = {
    'Individualios_pam': defaultdict(float),
    'Grupinės_pam': defaultdict(float)
}

# Žodynai su pamokų kainomis
individual_hourly_rates = {
    'Pamoka X': 30,
    'Pamoka Y': 20
}

group_hourly_rates = {
    'Mok_X, 11 kl.': 10,
    'Mok_Y, 11 kl.': 10,
    
    'Mok_X, 12 kl.': 10,
    'Mok_Y, 12 kl.': 10
}

# Žodynas su mokinių tėvų rekvizitais (kam išrašyti sąskaitą)
buyers = {
    'Pamoka X': 'Antanas Antanaitis \nPievų g. 1, LT-00000, Vilnius',
    'Pamoka Y': 'Petras Petraitis \nKlevų g. 5, LT-00001, Kaunas',

    'Mok_X, 11 kl.': 'Jonas Jonaitis \nMiškų g. 7, LT-00010, Vilnius',
    'Mok_Y, 11 kl.': 'Lukas Lukaitis \nMedžių g. 9, LT-00011, Klaipėda',

    'Mok_X, 12 kl.': 'Vardas Vardaitis \nBeržų g. 15, LT-00100, Kaunas',
    'Mok_Y, 12 kl.': 'Name Surname \nAdress str. 19, LT-00101, Klaipėda'
}

# Žodynas su mokinių tėvų emailais (kur išsiųsti sąskaitas)
name_to_email = {
    'Pamoka X': 'dominykdaunys@gmail.com',
    'Pamoka Y': 'dominykdaunys@gmail.com',

    'Mok_X, 11 kl.': 'dominykdaunys@gmail.com',
    'Mok_Y, 11 kl.': 'dominykdaunys@gmail.com',

    'Mok_X, 12 kl.': 'dominykdaunys@gmail.com',
    'Mok_Y, 12 kl.': 'dominykdaunys@gmail.com'
}

# Kalendorių, iš kurių imsime duomenis, pavadinimai
calendar_names = ['Individualios_pam', 'Grupinės_pam']

# Pradinis sąskaitos numeris
invoice_number = 3

# %% Suskaičiuoju, kiek valandų dirbau

# Einame per visus vartojo turimus kalendorius
for calendar in calendars:
    if calendar.name in calendar_names: # Jeigu atitinka dominantį kalendorių, tai vykdome...
        
        # Gauname įvykius nurodytam laikotarpiui
        try:
            events = calendar.date_search(start=start_date, end=end_date)
        except Exception as e:
            print(f"Error fetching events for calendar {calendar.name}: {e}")
            continue
        
        # Einame per kiekvieną įvykį iš kątik gauto laikotarpio
        for event in events:
            try:
                ical = Calendar.from_ical(event.data) # Sukuria eventui python objektą
            except Exception as e:
                print(f"Error parsing event data: {e}")
                continue     
            
            # Einame per gauto objekto (event) komponentus
            for component in ical.walk():
                if component.name == "VEVENT": # Mus domina tik ,,virtual events'', ignoruojame visa kitą
                    
                    # Išsitraukiame dominančią info
                    summary = str(component.get('summary'))
                    dtstart = component.get('dtstart').dt
                    dtend = component.get('dtend').dt
                    
                    # Apskaičiuojame trukmę AKADEMINĖMIS valandomis ir pridedame prie bendros sumos
                    if isinstance(dtstart, datetime) and isinstance(dtend, datetime):
                        duration = (dtend - dtstart).total_seconds() / 2700 # NES AKADEMINĖS VALANDOS
                        durations[calendar.name][summary] += duration

# %% Apskaičiuoju ir atvaizduoju kiek uždirbau iš kiekvieno mokinio

for calendar_name, activities in durations.items():
    print(f"\nProcessing calendar: {calendar_name}")
    for activity, total_hours in activities.items():
        if calendar_name == 'Individualios_pam' and activity in individual_hourly_rates:
            rate = individual_hourly_rates[activity]
        elif calendar_name == 'Grupinės_pam' and activity in group_hourly_rates:
            rate = group_hourly_rates[activity]
        else:
            rate = 0  # Default rate is 0 if not found
        earnings = total_hours * rate
        print(f"Activity: {activity}, Total Academic Hours: {total_hours:.2f}, Earnings: {earnings:.2f} EUR")

# %% Funkcija, kuri pakeičia šriftą

def set_font_to_times_new_roman(paragraph, size=12):
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.size = Pt(size)

# %% Kiekvieno mokinio sąskaitos generavimas

# Sukuriame sąrašą, kuriame saugome sugeneruotų sąskaitų pavadinimus
generated_invoices = []

for calendar_name, activities in durations.items():
    for activity, total_hours in activities.items():
        if calendar_name == 'Individualios_pam' and activity in individual_hourly_rates:
            rate = individual_hourly_rates[activity]
        elif calendar_name == 'Grupinės_pam' and activity in group_hourly_rates:
            rate = group_hourly_rates[activity]
        else:
            rate = 0  #Jeigu pamiršome nurodyti įkainį, bus nulis

        buyer_name = buyers.get(activity, "Unknown Buyer")
        total_amount = total_hours * rate

        # Užkrauname sąskaitos šabloną kiekvienam mokiniui
        doc = Document('pavyzdine_saskaita.docx')
        current_date = datetime.today()
        due_date = current_date + timedelta(weeks=1)
        amount_in_words = num2words.num2words(total_amount, lang='lt', to='currency')

        update_next_paragraph = False

        # Atnaujiname sąskaitos laukelius tekste
        for paragraph in doc.paragraphs:
            if update_next_paragraph:
                paragraph.text = buyer_name
                update_next_paragraph = False
            elif 'Serija ir Nr.' in paragraph.text:
                paragraph.text = f"Serija ir Nr. 2026-{invoice_number:02}"
            elif 'Sąskaitos data' in paragraph.text:
                paragraph.text = f"Sąskaitos data {current_date.strftime('%Y-%m-%d')}"
            elif 'Apmokėti iki' in paragraph.text:
                paragraph.text = f"Apmokėti iki {due_date.strftime('%Y-%m-%d')}"
            elif 'Suma žodžiais:' in paragraph.text:
                paragraph.text = f"Suma žodžiais: {amount_in_words.capitalize()}"
            elif 'Pirkėjas' in paragraph.text:
                update_next_paragraph = True
            
            set_font_to_times_new_roman(paragraph)  # Šriftas - Times new roman

        # Atnaujiname sąskaitos laukelius lentelėje
        for table in doc.tables:
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    if 'Kiekis' in cell.text and i + 1 < len(table.rows):
                        below_cell = table.cell(i + 1, j)
                        below_cell.text = str(int(total_hours))
                    elif 'Kaina' in cell.text and i + 1 < len(table.rows):
                        below_cell = table.cell(i + 1, j)
                        below_cell.text = f"{rate:.2f} €"
                    elif 'Iš viso' in cell.text and i + 1 < len(table.rows):
                        below_cell = table.cell(i + 1, j)
                        below_cell.text = f"{total_amount:.2f} €"
                    elif 'Bendra suma' in cell.text and j + 1 < len(row.cells):
                        next_cell = table.cell(i, j + 1)
                        next_cell.text = f"{total_amount:.2f} €"
                        
                    # Šriftas - Times new roman
                    for paragraph in cell.paragraphs:
                        set_font_to_times_new_roman(paragraph)

        # Išsaugome atnaujintą dokumentą
        output_filename = f'2026-{invoice_number:02}_{activity}.docx' # ŠITĄ REIKS KEIST KAS METAI
        output_path = os.path.abspath(f'Saskaitos_out/{output_filename}')
        doc.save(output_path)
        print(f"Invoice saved as {output_filename}")
        
        # Pridedu sąskaitos pavadinimą į sąrašą
        generated_invoices.append(f'2026-{invoice_number:02}_{activity}')

        # Sąskaitos numeris +1
        invoice_number += 1

# %% Funkcija, kuri pavercia docx i pdf

def docx_to_pdf(docx_path, pdf_path):
    # Load the Word application
    word = comtypes.client.CreateObject('Word.Application')
    # Minimize the Word application window
    word.Visible = False
    # Open the document in Word
    doc = word.Documents.Open(docx_path)
    # Save the document as a PDF
    doc.SaveAs(pdf_path, FileFormat=17)  # 17 is the file format for PDF in Word
    # Close the document and Word application
    doc.Close()
    word.Quit()

# %% Konvertuoju į pdf iš docx

# Einam per visas sugeneruotas sąskaitas, susikuriam nuorodas
for invoice_name in generated_invoices:
    docx_path = os.path.abspath(f'Saskaitos_out/{invoice_name}.docx')
    pdf_path = os.path.abspath(f'Saskaitos_out/{invoice_name}.pdf')

    # Tikrinu, ar docx failas egzistuoja
    if os.path.exists(docx_path):
        try:
            # Konvertuojam docx į pdf
            docx_to_pdf(docx_path, pdf_path)
            print(f"Converted to PDF: {pdf_path}")
        except Exception as e:
            print(f"Failed to convert {docx_path} to PDF: {e}")
    else:
        print(f"File does not exist: {docx_path}")

# %% Pervadinu failus

# Kelias iki aplanko su sąskaitomis
invoices_dir = os.path.abspath('Saskaitos_out')

# Pervadinu pdf failus, pašalindamas viską, kas po ,,_'' ženklo. 
# Pirmiausia sukuriu naują pavadinimą ir kelius iki failo
for invoice_name in generated_invoices:
    new_name = invoice_name.split('_')[0] + '.pdf'
    old_path = os.path.join(invoices_dir, invoice_name + '.pdf')
    new_path = os.path.join(invoices_dir, new_name)

    # Pervadinu failus
    if os.path.exists(old_path):
        os.rename(old_path, new_path)
        print(f'Renamed {old_path} to {new_path}')
    else:
        print(f'File not found for renaming: {old_path}')

# %% Susieju gavėjų adresus su atitinkamu failu

recipient_info = {}

for invoice_name in generated_invoices:
    # Ištraukiame mokinio vardą iš sąskaitos failo pavadinimo
    name_part = "_".join(invoice_name.split('_')[1:])
    
    # Gauname atitinkamą el. pašto adresą
    recipient_email = name_to_email.get(name_part)
    
    # Pridedame į recipient_info, jei el. pašto adresas buvo rastas
    if recipient_email:
        recipient_info[invoice_name.split('_')[0]] = recipient_email

# Patikriname rezultatą
print(recipient_info)

# %% Funkcija, kuri išsiunčia emailą

def send_email(recipient_email, subject, body, attachment_path):
    # Create a multipart message
    msg = MIMEMultipart()
    msg['From'] = gmail_user
    msg['To'] = recipient_email
    msg['Subject'] = subject

    # Attach the body of the email
    msg.attach(MIMEText(body, 'html'))

    # Attach the file
    attachment_name = os.path.basename(attachment_path)
    attachment = MIMEBase('application', 'octet-stream')
    with open(attachment_path, 'rb') as attachment_file:
        attachment.set_payload(attachment_file.read())
    encoders.encode_base64(attachment)
    attachment.add_header('Content-Disposition', f'attachment; filename={attachment_name}')
    msg.attach(attachment)

    # Connect to Gmail server and send the email
    try:
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(gmail_user, gmail_password)
        server.sendmail(gmail_user, recipient_email, msg.as_string())
        server.close()
        print(f"Email sent successfully to {recipient_email}")
    except Exception as e:
        print(f"Failed to send email to {recipient_email}. Error: {str(e)}")

# %% ŠIS KODO BLOKAS IŠSIUNČIA LAIŠKUS! PATIKRINKITE, AR TIKRAI NORITE TAI PADARYTI!

# Prašome patvirtinimo prieš siunčiant sąskaitas
confirmation = input("Did you check the month? Are you sure you want to send all invoices? (yes/no): ").strip().lower()

if confirmation == 'yes':
    # Einame per sugeneruotas sąskaitas ir jas išsiunčiame
    for invoice_prefix, recipient_email in recipient_info.items():
        pdf_path = os.path.join(invoices_dir, invoice_prefix + '.pdf')
        
        # Prieš siunčiant patikriname, ar pdf'as egzistuoja
        if os.path.exists(pdf_path):
            # email_subject = f'Sąskaita už pamokas ({datetime.now().strftime("%B")})' #DABARTINIS MENESIS
            last_month = (datetime.now().replace(day=1) - timedelta(days=1)).strftime("%B")  #PRAEJES MENESIS
            email_subject = f'Sąskaita už pamokas ({last_month})' #PRAEJES MENESIS
            
            # Sukuriame HTML laiško tekstą
            email_body = f"""
            <p>Laba diena.</p>
            <p>Siunčiu sąskaitą už pamokas.</p>
            <p>Pagarbiai<br>
            <span style="color: grey;">Dominykas Daunys</span></p>
            """
            
            send_email(recipient_email, email_subject, email_body, pdf_path)
        else:
            print(f"Invoice PDF not found: {pdf_path}")

else:
    print('Sending invoices has been canceled.')